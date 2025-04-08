#!/usr/bin/env python3

import json
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from typing import Dict, List, Any
from pathlib import Path
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.enums import TA_LEFT, TA_CENTER

def load_and_process_data(file_path: str) -> pd.DataFrame:
    """Load exam results from JSON and convert to DataFrame.
    
    Args:
        file_path (str): Path to the JSON file containing exam results
        
    Returns:
        pd.DataFrame: Processed exam results data
    """
    try:
        with open(file_path, 'r') as f:
            data = json.load(f)
        df = pd.DataFrame(data)
        return df
    except Exception as e:
        print(f"Error loading data: {str(e)}")
        return pd.DataFrame()

def calculate_student_performance(df: pd.DataFrame, student_id: int) -> Dict[str, Any]:
    """Calculate performance metrics for a specific student compared to their cohort.
    
    Args:
        df (pd.DataFrame): The complete exam results dataset
        student_id (int): ID of the student to analyze
        
    Returns:
        Dict[str, Any]: Dictionary containing performance metrics
    """
    # Filter for the specific student
    student_data = df[df['studentId'] == student_id]
    student_cohort = student_data['cohort'].iloc[0]
    
    # Get cohort data
    cohort_data = df[df['cohort'] == student_cohort]
    
    # Calculate metrics
    metrics = {
        'student_id': student_id,
        'cohort': student_cohort,
        'calendar_year': student_data['calendarYear'].iloc[0],
        'overall_average': student_data['responseValue'].mean(),
        'cohort_average': cohort_data['responseValue'].mean(),
        'performance_by_group': {},
        'performance_by_teaching_period': {}
    }
    
    # Calculate performance by item group
    for group in df['itemGroupCode'].unique():
        student_group_avg = student_data[student_data['itemGroupCode'] == group]['responseValue'].mean()
        cohort_group_avg = cohort_data[cohort_data['itemGroupCode'] == group]['responseValue'].mean()
        metrics['performance_by_group'][group] = {
            'student_avg': student_group_avg,
            'cohort_avg': cohort_group_avg
        }
    
    # Calculate performance by teaching period
    for period in df['teachingPeriod'].unique():
        student_period_avg = student_data[student_data['teachingPeriod'] == period]['responseValue'].mean()
        cohort_period_avg = cohort_data[cohort_data['teachingPeriod'] == period]['responseValue'].mean()
        metrics['performance_by_teaching_period'][period] = {
            'student_avg': student_period_avg,
            'cohort_avg': cohort_period_avg
        }
    
    return metrics

def generate_performance_plots(metrics: Dict[str, Any], output_dir: str) -> List[str]:
    """Generate performance comparison plots.
    
    Args:
        metrics (Dict[str, Any]): Performance metrics dictionary
        output_dir (str): Directory to save the plots
        
    Returns:
        List[str]: List of generated plot file paths
    """
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)
    plot_files = []
    
    # Plot overall performance comparison
    plt.figure(figsize=(10, 6))
    plt.bar(['Student', 'Cohort'], 
            [metrics['overall_average'], metrics['cohort_average']],
            color=['skyblue', 'lightgray'])
    plt.title('Overall Performance Comparison')
    plt.ylabel('Average Score')
    overall_plot = output_path / 'overall_performance.png'
    plt.savefig(overall_plot)
    plt.close()
    plot_files.append(str(overall_plot))
    
    # Plot group performance
    group_data = pd.DataFrame(metrics['performance_by_group']).T
    plt.figure(figsize=(12, 6))
    group_data.plot(kind='bar')
    plt.title('Performance by Item Group')
    plt.xlabel('Item Group')
    plt.ylabel('Average Score')
    plt.xticks(rotation=45)
    plt.tight_layout()
    group_plot = output_path / 'group_performance.png'
    plt.savefig(group_plot)
    plt.close()
    plot_files.append(str(group_plot))
    
    return plot_files

def main():
    # Example usage
    file_path = "jsData.json"
    df = load_and_process_data(file_path)
    
    if not df.empty:
        # Get unique student IDs
        student_ids = df['studentId'].unique()
        
        # Process first student as an example
        if len(student_ids) > 0:
            student_metrics = calculate_student_performance(df, student_ids[0])
            plots = generate_performance_plots(student_metrics, 'output')
            
            print(f"Processed performance report for student {student_ids[0]}")
            print(f"Overall average: {student_metrics['overall_average']:.2f}")
            print(f"Cohort average: {student_metrics['cohort_average']:.2f}")
            print(f"Plots generated: {plots}")

if __name__ == "__main__":
    main()
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_introduction_section(student_id: str) -> List[Any]:
    """Create the introduction section of the report.
    
    Args:
        student_id (str): The student's ID number
        
    Returns:
        List[Any]: List of flowable elements for the PDF
    """
    styles = getSampleStyleSheet()
    custom_style = ParagraphStyle(
        'CustomStyle',
        parent=styles['Normal'],
        spaceBefore=12,
        spaceAfter=12,
        leading=16
    )
    
    # Create elements list
    elements = []
    
    # Add student ID
    student_id_text = Paragraph(f"Student's ID: {student_id}", styles['Heading1'])
    elements.append(student_id_text)
    elements.append(Spacer(1, 12))
    
    # Add Introduction heading
    intro_heading = Paragraph("Introduction", styles['Heading2'])
    elements.append(intro_heading)
    elements.append(Spacer(1, 12))
    
    # Add introduction text
    intro_text = Paragraph(
        "To help you interpret the information and tables of data in this feedback please note:",
        custom_style
    )
    elements.append(intro_text)
    
    # Add bullet points
    bullet_points = [
        "The score is the number of correct answers you achieved at SBAs.",
        "All scores were then converted into Action University's Generic Marking Scale (AUGMS) for undergraduate programmes, with pass-mark fixed at 35.50%.",
        "A pass-score was determined by means of a standard setting methodology, equivalent to the 35.50% pass-mark in AUGMS. So to find out if you have achieved a pass, please check if your overall mark is equal to or over 35.50%.",
        "It may help you understand what question format (SBA) you need to work at further, by looking at your raw scores for each format. However, please note that in this exam there is considerable compensation across question formats, and in any one exam you will only need to achieve an overall pass for question formats combined.",
        "You will also want to look at your performance in each domain: since this gives some idea of your relative strengths. However, when the total number of questions in each domain is low, this will not be very accurate, so please use this information as a guide only.",
        "When you have low scores in a number of domains it usually means you need to study all domains more thoroughly - this is an important take-home message."
    ]
    
    bullet_list = ListFlowable(
        [ListItem(Paragraph(point, custom_style)) for point in bullet_points],
        bulletType='bullet',
        start='bulletchar',
        bulletFontSize=8,
        leftIndent=20,
        bulletOffsetY=2
    )
    elements.append(bullet_list)
    
    return elements

def generate_student_report(json_data: Dict[str, Any], output_filename: str) -> None:
    """Generate a PDF report for a student.
    
    Args:
        json_data (Dict[str, Any]): The student's data
        output_filename (str): The name of the output PDF file
    """
    # Create the PDF document
    doc = SimpleDocTemplate(
        output_filename,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Create the elements list
    elements = []
    
    # Add introduction section
    elements.extend(create_introduction_section(str(json_data['student_id'])))
    
    # Build the PDF
    doc.build(elements)
    # Create a new Word document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = 12

    # Add student ID
    doc.add_paragraph(f"Student's ID: {json_data['student_id']}")

    # Add Introduction section
    doc.add_heading('Introduction', level=2)
    intro_text = """
To help you interpret the information and tables of data in this
feedback please note:

-   The score is the number of correct answers you achieved at SBAs.

-   All scores were then converted into Aston University's Generic
    Marking Scale (AUGMS) for undergraduate programmes, with pass-mark
    fixed at 39.50%.

-   A pass-score was determined by means of a standard setting
    methodology, equivalent to the 39.50% pass-mark in AUGMS. So to find
    out if you have achieved a pass, please check if your overall mark
    is equal to or over 39.50%.

-   It may help you understand what question format (SBA) you need to
    work at further, by looking at your raw scores for each format.
    However bear in mind that in this exam there is complete
    compensation across question formats, and at any one exam you will
    only need to achieve an overall pass for question formats combined.

-   You will also want to look at your performance in each domain, since
    this gives some idea of your relative strengths. However, when the
    total number of questions in each domain is low, this will not be
    very accurate, so please use this information as a guide only.

-   When you have low scores in a number of domains it usually means you
    need to study all domains more thoroughly - this is an important
    take-home message.
"""
    doc.add_paragraph(intro_text)
    doc.add_page_break()

    # Add Overall Performance section
    doc.add_heading('Your overall performance (MARKS)', level=2)
    doc.add_paragraph("""
The graph below shows your maximum percentage MARK for all question
formats combined and the overall pass-mark fixed at 39.50% (rounded to
40% and equivalent to the overall pass-score obtained via the Angoff
standard setting method, equal to 56.8 out of 100).
""")

    # Generate overall performance chart
    plt.figure(figsize=(8, 6))
    student_score = json_data['summary_results'][0]['your_score']
    pass_mark = 39.5
    max_score = 100

    bars = plt.bar(['Your Score', 'Pass Mark', 'Max Score'],
                  [student_score, pass_mark, max_score],
                  color=['#1f77b4', '#ff7f0e', '#2ca02c'])

    plt.ylabel('Percentage (%)')
    plt.title('Overall Performance Comparison')
    plt.ylim(0, 110)

    # Add value labels on top of bars
    for bar in bars:
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2., height,
                 f'{height}%',
                 ha='center', va='bottom')

    # Save chart to memory
    chart1 = BytesIO()
    plt.savefig(chart1, format='png', dpi=300)
    plt.close()

    # Add chart to document
    chart1.seek(0)
    doc.add_picture(chart1, width=Inches(5.15))

    # Add overall outcome
    outcome_para = doc.add_paragraph()
    outcome_para.add_run(f"Your Overall Outcome at the Current Exam: {json_data['overall_outcome']}").italic = True
    outcome_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Add Mark Range table
    doc.add_paragraph("""
Please note: If your outcome was Borderline Pass, that means that you
have passed the current exam, but it should draw your attention to how
close your performance was to the pass-score. Grades and Descriptors
listed are for indicative purposes only. Please note your end of year
transcript for summative exams will only list marks and decisions.
""")

    # Create mark range table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Mark Range (Lower Bound)'
    hdr_cells[1].text = 'Mark Range (Upper Bound)'
    hdr_cells[2].text = 'Descriptor'
    hdr_cells[3].text = 'Grade'

    # Add data rows
    for range_data in json_data['mark_ranges']:
        row_cells = table.add_row().cells
        row_cells[0].text = range_data['lower']
        row_cells[1].text = range_data['upper']
        row_cells[2].text = range_data['descriptor']
        row_cells[3].text = range_data['grade']

    doc.add_page_break()

    # Add Summary of Results section
    doc.add_heading('Your summary of results (RAW SCORES)', level=2)
    doc.add_paragraph("""
The table below summarises your SCORES for the overall exam (%), SBAs.
The class min, max, mean and standard deviation of scores are also
displayed.
""")

    # Create summary table
    summary_table = doc.add_table(rows=1, cols=7)
    summary_table.style = 'Table Grid'

    # Header row
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = 'Exam component'
    hdr_cells[1].text = 'Your score'
    hdr_cells[2].text = 'Total Available'
    hdr_cells[3].text = 'Min'
    hdr_cells[4].text = 'Max'
    hdr_cells[5].text = 'Mean'
    hdr_cells[6].text = 'StDev'

    # Add data rows
    for component in json_data['summary_results']:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = component['component']
        row_cells[1].text = str(component['your_score'])
        row_cells[2].text = str(component['total_available'])
        row_cells[3].text = str(component['min'])
        row_cells[4].text = str(component['max'])
        row_cells[5].text = str(component['mean'])
        row_cells[6].text = str(component['stdev'])

    doc.add_page_break()

    # Add Decile Ranking section
    doc.add_heading('Decile ranking', level=2)
    doc.add_paragraph("""
Your decile rank is reported below. This is a measure of your
performance in comparison to the performance of the other students who
sat the exam.
""")

    decile_para = doc.add_paragraph()
    decile_para.add_run(f"You are in the {json_data['decile_rank']} decile.").italic = True
    decile_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("How to interpret your rank:")

    decile_explanation = [
        "1st decile: top 10% of the cohort",
        "2nd decile: top 20% of the cohort",
        "3rd decile: top 30% of the cohort",
        "4th decile: top 40% of the cohort",
        "5th decile: top 50% of the cohort",
        "6th decile: bottom 50% of the cohort",
        "7th decile: bottom 40% of the cohort",
        "8th decile: bottom 30% of the cohort",
        "9th decile: bottom 20% of the cohort",
        "10th decile: bottom 10% of the cohort"
    ]

    for item in decile_explanation:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("""
In reading this information, please bear in mind:

-   The decile rank here presented does NOT imply any pass / fail
    decision nor does it correspond to any prize or merit.

-   This decile rank does NOT translate directly into the ranking for
    the UK Foundation Programme, the latter being the result of
    weighting across all summative exams in the MBChB, from Year 1 to
    Year 5.

-   Students may leave or join your cohort later in the programme, and
    this could shift the final ranking for the UK Foundation Programme.

For all the reasons listed above, please only use this information
formatively, aiming to inform and improve your future learning.
""")

    doc.add_page_break()

    # Add SBA Results section
    doc.add_heading('Your results for Single Best Answer (SBA) Questions', level=2)
    doc.add_paragraph("""
The graph below shows your SCORE for the SBAs and your performance. The
table below the graph displays the number of questions you got correct,
the total number of questions available, and the minimum, maximum, and
mean score computed at the cohort's level, organised by block. Please
bear in mind that in this exam there is complete compensation across
question formats, and you only need to achieve an overall pass for
question formats combined.
""")

    # Generate SBA performance chart
    plt.figure(figsize=(10, 6))
    blocks = [block['block'] for block in json_data['sba_results']]
    your_scores = [block['your_score'] for block in json_data['sba_results']]
    max_scores = [block['total'] for block in json_data['sba_results']]
    cohort_means = [block['mean'] for block in json_data['sba_results']]

    x = range(len(blocks))
    width = 0.25

    plt.bar([i - width for i in x], your_scores, width, label='Your Score', color='#1f77b4')
    plt.bar(x, cohort_means, width, label='Cohort Mean', color='#ff7f0e')
    plt.bar([i + width for i in x], max_scores, width, label='Max Score', color='#2ca02c')

    plt.xlabel('Blocks')
    plt.ylabel('Scores')
    plt.title('SBA Performance by Block')
    plt.xticks(x, blocks, rotation=45, ha='right')
    plt.legend()
    plt.tight_layout()

    # Save chart to memory
    chart2 = BytesIO()
    plt.savefig(chart2, format='png', dpi=300)
    plt.close()

    # Add chart to document
    chart2.seek(0)
    doc.add_picture(chart2, width=Inches(6.5))

    doc.add_page_break()

    # Create SBA results table
    sba_table = doc.add_table(rows=1, cols=7)
    sba_table.style = 'Table Grid'

    # Header row
    hdr_cells = sba_table.rows[0].cells
    hdr_cells[0].text = 'Block'
    hdr_cells[1].text = 'Your Score'
    hdr_cells[2].text = 'Total'
    hdr_cells[3].text = 'Min'
    hdr_cells[4].text = 'Max'
    hdr_cells[5].text = 'Mean'
    hdr_cells[6].text = 'StDev'

    # Add data rows
    for block in json_data['sba_results']:
        row_cells = sba_table.add_row().cells
        row_cells[0].text = block['block']
        row_cells[1].text = str(block['your_score'])
        row_cells[2].text = str(block['total'])
        row_cells[3].text = str(block['min'])
        row_cells[4].text = str(block['max'])
        row_cells[5].text = str(block['mean'])
        row_cells[6].text = str(block['stdev'])

    # Save the document
    doc.save(output_filename)

# Example JSON data for one student
student_data = {
    "student_id": "12345678",
    "overall_outcome": "Very Good Pass",
    "mark_ranges": [
        {"lower": "69.50%", "upper": ">", "descriptor": "Excellent Pass", "grade": "A"},
        {"lower": "59.50%", "upper": "69.49%", "descriptor": "Very Good Pass", "grade": "B"},
        {"lower": "49.50%", "upper": "59.49%", "descriptor": "Good Pass", "grade": "C"},
        {"lower": "44.50%", "upper": "49.49%", "descriptor": "Pass", "grade": "D1"},
        {"lower": "39.50%", "upper": "44.49%", "descriptor": "Borderline Pass", "grade": "D2"},
        {"lower": "<", "upper": "39.49%", "descriptor": "NOT Pass", "grade": "E"}
    ],
    "summary_results": [
        {"component": "Overall Scores", "your_score": 68, "total_available": 100, "min": 44, "max": 93, "mean": 70.49, "stdev": 10.04},
        {"component": "SBAs", "your_score": 68, "total_available": 100, "min": 44, "max": 93, "mean": 70.00, "stdev": 10.04}
    ],
    "decile_rank": "7th",
    "sba_results": [
        {"block": "Cardiovascular System", "your_score": 7, "total": 16, "min": 2, "max": 14, "mean": 9.48, "stdev": 2.47},
        {"block": "Cell Biology & Genetics", "your_score": 2, "total": 2, "min": 0, "max": 2, "mean": 1.79, "stdev": 0.45},
        {"block": "ICP - Clinical Skills", "your_score": 6, "total": 6, "min": 3, "max": 6, "mean": 4.95, "stdev": 0.88},
        {"block": "ICP - Ethics & Law", "your_score": 1, "total": 2, "min": 0, "max": 2, "mean": 1.28, "stdev": 0.61},
        {"block": "Infection", "your_score": 12, "total": 17, "min": 7, "max": 17, "mean": 13.89, "stdev": 2.49},
        {"block": "Metabolism, Endocrinology & Haematology", "your_score": 12, "total": 16, "min": 6, "max": 15, "mean": 11.84, "stdev": 2.11},
        {"block": "Musculoskeletal System", "your_score": 15, "total": 18, "min": 4, "max": 18, "mean": 12.15, "stdev": 2.95},
        {"block": "Pathological Processes", "your_score": 8, "total": 15, "min": 5, "max": 15, "mean": 10.48, "stdev": 2.22},
        {"block": "Physiology & Pharmacology", "your_score": 2, "total": 5, "min": 0, "max": 5, "mean": 2.72, "stdev": 1.10},
        {"block": "Social & Psychological Aspects of Health", "your_score": 3, "total": 3, "min": 0, "max": 3, "mean": 1.93, "stdev": 0.79}
    ]
}

# First install required packages
!pip install python-docx matplotlib

# Generate the report
generate_student_report(student_data, "Student_Report.docx")

print("Report generated successfully as Student_Report.docx")